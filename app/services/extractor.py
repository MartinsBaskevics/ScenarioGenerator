import re

from docx import Document
from docx.oxml.ns import qn

from app.models import Requirement

# atpazīst funkcionālās prasības virsrakstu, lai saprastu kur tās sākas
FP_SECTION_START = re.compile(r"\bfunkcionāl[aā][s]?\s*prasīb", re.IGNORECASE)

# atpazīst citus virsrakstus, lai saprastu kur beidzas funkcionālās prasības
FP_SECTION_END = re.compile(
    r"nefunkcionāl|pielikum|izmantot.*avot|rezultāt[i]?\s*un|"
    r"secinājum|drošīb.*prasīb|uzturēšan|darbu\s+gaita|nodevum",
    re.IGNORECASE,
)

# atpazīst FP-1, MN-FP-01, KV-FP-01 u.c. formātus
FP_ID_PATTERN = re.compile(
    r"^((?:[A-ZĒŪĪĀŠĢĶĻŽČŅ]{1,4}-)?FP-\d+)\s*[:\s]*(.*)",
    re.IGNORECASE,
)

# atpazīst 1.1. 1.2. 1.3. formātus
NUMBERED_PATTERN = re.compile(r"^(\d+\.\d+\.\d+(?:\.\d+){0,2})[.\s]+(.*)")

# atpazīst tabulu kolonnu nosaukumus, lai zinātu no kuras jāņem prasības
TABLE_HEADER_PATTERN = re.compile(
    r"^(prasīb[au]s?\s*(id|apraksts|nosaukums)|prasīb[au])$",
    re.IGNORECASE,
)

# atpazīst darbības vārdus, lai noteiktu vai ir prasība vai virsraksts
ACTION_WORDS = re.compile(
    r"\b(nodrošināt|jānodrošina|iespēja|jāparedz|veikt|atbalstīt|"
    r"rādīt|kārtot|dzēst|mainīt|pievienot|ievietot|iegūt|saglabāt|"
    r"eksportēt|sniegt|ģenerēt|pārvaldīt|integrēt|uzturēt|atļaut|"
    r"ierobežot|parādīt|izveidot|definēt|norādīt|reģistrēt)\b",
    re.IGNORECASE,
)


def normalizeId(raw):
    return raw.strip().rstrip(".:")


def normalizeText(raw):
    return re.sub(r"\s+", " ", raw).strip()


# atšķir virsrakstus
def looksLikeHeading(text):
    # garš teksts nevar būt virsraksts
    if len(text) > 150:
        return False
    if text == text.upper() and len(text) > 3:
        return True
    if text.endswith(":") and re.match(r"^[A-ZĒŪĪĀŠĢĶĻŽČŅ]", text):
        return True
    if re.match(r"^#{1,4}\s", text):
        return True
    if re.match(r"^\d+(?:\.\d+)*\.\s+\S", text):
        return True
    return False


def getSectionNumber(text):
    m = re.match(r"^(\d+)[.\s]", text.strip())
    return m.group(1) if m else None


def isSectionEnd(line, currentSectionNum):
    if not line:
        return False
    if FP_SECTION_END.search(line) and looksLikeHeading(line):
        return True
    if currentSectionNum:
        top = getSectionNumber(line)
        if top and top != currentSectionNum and looksLikeHeading(line):
            return True
    return False


def isNumberedHeading(text):
    # ja satur darbības vārdu, tā ir prasība nevis virsraksts
    if ACTION_WORDS.search(text):
        return False
    if text.endswith(":"):
        return True
    if len(text) < 65 and not text.endswith(";") and re.match(r"^[A-ZĒŪĪĀŠĢĶĻŽČŅA-Z]", text):
        return True
    return False


def getWordListLevel(paragraph):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    el = numPr.find(qn("w:ilvl"))
    return int(el.get(qn("w:val"))) if el is not None else None


# meklē pēc FP-1, MN-FP-01 utt.
def extractByFpIds(lines):
    found = []
    current = None
    for line in lines:
        line = line.strip()
        if not line or TABLE_HEADER_PATTERN.match(line):
            continue
        m = FP_ID_PATTERN.match(line)
        if m:
            if current:
                found.append(current)
            current = Requirement(id=normalizeId(m.group(1)), text=normalizeText(m.group(2)))
        elif current:
            extra = normalizeText(line)
            if extra and len(extra) > 5 and not looksLikeHeading(extra):
                current.text = (current.text + " " + extra).strip()
    if current:
        found.append(current)
    return found


# meklē pēc 4.1.1, 4.2.3 utt.
def extractByNumbers(lines):
    found = []
    current = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        m = NUMBERED_PATTERN.match(line)
        if m:
            reqId = m.group(1)
            reqText = normalizeText(m.group(2))
            depth = reqId.count(".") + 1

            if depth < 3:
                continue

            if isNumberedHeading(reqText):
                continue

            parentId = ".".join(reqId.split(".")[:-1])
            isChild = current is not None and parentId == current.id

            if isChild:
                if reqText:
                    current.text = (current.text + " " + reqText).strip()
            else:
                if current:
                    found.append(current)
                current = Requirement(id=reqId, text=reqText)

        elif current and len(line) > 5 and not looksLikeHeading(line):
            current.text = (current.text + " " + normalizeText(line)).strip()

    if current:
        found.append(current)
    return found


# izvēlas stratēģiju apstrādei
def autoDetectAndExtract(fpLines):
    hasFpIds = any(FP_ID_PATTERN.match(l.strip()) for l in fpLines)
    hasNumbers = any(NUMBERED_PATTERN.match(l.strip()) for l in fpLines)

    if hasFpIds:
        return extractByFpIds(fpLines)
    if hasNumbers:
        return extractByNumbers(fpLines)

    return [
        Requirement(id=f"FP-{i + 1}", text=normalizeText(line.strip()))
        for i, line in enumerate(fpLines)
        if len(line.strip()) > 30
        and not looksLikeHeading(line.strip())
        and not TABLE_HEADER_PATTERN.match(line.strip())
    ]


# sadala tekstu rindās, atrod FP sadaļu un izsauc prasības izguvēju
def fromText(text):
    lines = text.splitlines()
    fpLines = []
    inFpSection = False
    sectionNum = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if not inFpSection:
            if FP_SECTION_START.search(stripped):
                inFpSection = True
                sectionNum = getSectionNumber(stripped)
            continue
        if isSectionEnd(stripped, sectionNum):
            inFpSection = False
            sectionNum = None
            continue
        if FP_SECTION_START.search(stripped):
            continue
        fpLines.append(line)

    return autoDetectAndExtract(fpLines)

# atver Word dokumentu un izlemj, kādu FP izguvēju lietot
def fromDocx(path):
    doc = Document(path)

    tableReqs = extractFromDocxTables(doc)
    if tableReqs:
        return tableReqs

    if any(getWordListLevel(p) is not None for p in doc.paragraphs):
        return extractByWordListLevel(doc)

    return extractFromDocxParagraphs(doc)


def iterDocxBlocks(doc):
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            yield "para", text
        elif tag == "tbl":
            yield "table", child


def getTableRows(tblElem):
    rows = []
    for tr in tblElem.iter(qn("w:tr")):
        cells = []
        for tc in tr.iter(qn("w:tc")):
            text = "".join(t.text or "" for t in tc.iter(qn("w:t")))
            cells.append(text.strip())
        if cells:
            rows.append(cells)
    return rows


# tabulu apstrādei
def extractFromDocxTables(doc):
    found = []
    inFpSection = False
    sectionNum = None

    for kind, item in iterDocxBlocks(doc):
        if kind == "para":
            text = normalizeText(item)
            if not text:
                continue
            if not inFpSection:
                if FP_SECTION_START.search(text):
                    inFpSection = True
                    sectionNum = getSectionNumber(text)
                continue
            if isSectionEnd(text, sectionNum):
                inFpSection = False
                sectionNum = None
                continue
            if FP_SECTION_START.search(text):
                continue

        elif kind == "table":
            if not inFpSection:
                continue
            rows = getTableRows(item)
            current = None
            for row in rows:
                if len(row) < 2:
                    continue
                first = row[0].strip()
                if TABLE_HEADER_PATTERN.match(first):
                    continue
                m = FP_ID_PATTERN.match(first)
                if m:
                    if current:
                        found.append(current)
                    desc = row[2] if len(row) >= 3 and row[2] else row[1]
                    current = Requirement(id=normalizeId(m.group(1)), text=normalizeText(desc))
                elif current:
                    extra = normalizeText(" ".join(row))
                    if extra and len(extra) > 5 and not TABLE_HEADER_PATTERN.match(extra):
                        current.text = (current.text + " " + extra).strip()
            if current:
                found.append(current)

    return found


# Word automātiskā numurēšana - skaitļi ir XML ilvl atribūtā, nevis tekstā
# ilvl 0,1 ir sadaļu virsraksti, ilvl 2+ ir prasības
def extractByWordListLevel(doc):
    inFpSection = False
    sectionNum = None
    found = []
    current = None
    reqLevel = None

    for p in doc.paragraphs:
        text = normalizeText(p.text)
        lvl = getWordListLevel(p)
        if not text:
            continue

        if not inFpSection:
            if FP_SECTION_START.search(text):
                inFpSection = True
                sectionNum = getSectionNumber(text)
            continue

        if isSectionEnd(text, sectionNum):
            inFpSection = False
            sectionNum = None
            if current:
                found.append(current)
                current = None
            reqLevel = None
            continue

        if FP_SECTION_START.search(text):
            continue

        if lvl is None:
            if current and len(text) > 5:
                current.text = (current.text + " " + text).strip()
            continue

        if lvl < 2:
            continue

        if isNumberedHeading(text):
            if current is None:
                reqLevel = lvl + 1
            elif lvl > reqLevel:
                pass
            else:
                found.append(current)
                current = None
                reqLevel = lvl + 1
            continue

        if reqLevel is None:
            reqLevel = lvl

        if lvl == reqLevel:
            if current:
                found.append(current)
            current = Requirement(id=f"FP-{len(found) + 1}", text=text)
        elif lvl > reqLevel:
            if current:
                current.text = (current.text + " " + text).strip()
        else:
            if current:
                found.append(current)
            reqLevel = lvl
            current = Requirement(id=f"FP-{len(found) + 1}", text=text)

    if current:
        found.append(current)
    return found


def extractFromDocxParagraphs(doc):
    fpLines = []
    inFpSection = False
    sectionNum = None

    for p in doc.paragraphs:
        line = normalizeText(p.text)
        if not line:
            continue
        if not inFpSection:
            if FP_SECTION_START.search(line):
                inFpSection = True
                sectionNum = getSectionNumber(line)
            continue
        if isSectionEnd(line, sectionNum):
            inFpSection = False
            sectionNum = None
            continue
        if FP_SECTION_START.search(line):
            continue
        fpLines.append(line)

    return autoDetectAndExtract(fpLines)
